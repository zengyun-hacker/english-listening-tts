"""
Export a ListeningTest to a Word (.docx) document.

Structure:
  - Title (Heading 1)
  - For each section:
      - Section name (Heading 2)
      - Instructions (italic paragraph)
      - For each item:
          - Question number + question text (Bold)
          - Options A / B / C / D
  - Answer Key page (Heading 2)
      - Answers with explanations grouped by section
"""

from __future__ import annotations

import io
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from llm.deepseek import ListeningTest


def export_to_docx(test: "ListeningTest") -> bytes:
    """
    Generate a Word document from a ListeningTest and return it as bytes.

    Args:
        test: A populated ListeningTest dataclass.

    Returns:
        Raw bytes of the .docx file, ready for st.download_button().
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # ── Page margins (narrower for more content) ──────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── Title ─────────────────────────────────────────────────────────────
    title_para = doc.add_heading(test.title, level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()   # spacer

    # ── Sections ──────────────────────────────────────────────────────────
    for sec in test.sections:
        doc.add_heading(sec.name, level=2)

        if sec.instructions:
            instr = doc.add_paragraph(sec.instructions)
            instr.runs[0].italic = True
            instr.paragraph_format.space_after = Pt(6)

        for item in sec.items:
            # Question header: "1. What does the woman imply?"
            q_para = doc.add_paragraph()
            q_run = q_para.add_run(f"{item.number}. {item.question}")
            q_run.bold = True
            q_para.paragraph_format.space_before = Pt(8)
            q_para.paragraph_format.space_after  = Pt(2)

            # Options
            for letter in ("A", "B", "C", "D"):
                opt_text = item.options.get(letter, "")
                if opt_text:
                    opt_para = doc.add_paragraph(style="List Bullet")
                    opt_para.clear()   # remove auto bullet
                    opt_run = opt_para.add_run(f"  {letter}. {opt_text}")
                    opt_para.paragraph_format.space_after = Pt(1)

        doc.add_paragraph()   # spacer between sections

    # ── Answer Key ────────────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("Answer Key  参考答案", level=2)

    for sec in test.sections:
        sec_heading = doc.add_paragraph()
        sec_run = sec_heading.add_run(f"【{sec.name}】")
        sec_run.bold = True
        sec_run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

        for item in sec.items:
            ans_para = doc.add_paragraph()
            ans_para.paragraph_format.left_indent = Inches(0.3)
            ans_run = ans_para.add_run(f"{item.number}. {item.answer}")
            ans_run.bold = True
            if item.explanation:
                ans_para.add_run(f"   {item.explanation}")

        doc.add_paragraph()   # spacer

    # ── Serialize to bytes ────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_to_docx_full(test: "ListeningTest") -> bytes:
    """
    Generate a Word document containing script + questions + answers for each item.
    Useful as a teacher's reference or post-exam review copy.

    Structure per item:
        Number X.
        [Script — shaded block]
        Question + Options
        ✓ Answer + Explanation
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def _set_cell_bg(cell, hex_color: str):
        """Set paragraph shading (simulated via run highlight isn't enough; use XML)."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    def _add_shaded_script(doc, script_text: str):
        """Add script inside a 1-cell table with grey background."""
        table = doc.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        cell = table.cell(0, 0)
        _set_cell_bg(cell, "F3F4F6")   # light grey

        # Clear default empty paragraph in cell
        cell.paragraphs[0].clear()
        lines = script_text.strip().splitlines()
        for i, line in enumerate(lines):
            if i == 0:
                p = cell.paragraphs[0]
            else:
                p = cell.add_paragraph()
            run = p.add_run(line)
            run.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x37, 0x47, 0x51)
            p.paragraph_format.space_after = Pt(1)

        # spacer after table
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── Title ─────────────────────────────────────────────────────────────
    title_para = doc.add_heading(f"{test.title}  【含原文版】", level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # ── Sections ──────────────────────────────────────────────────────────
    for sec in test.sections:
        doc.add_heading(sec.name, level=2)
        if sec.instructions:
            instr = doc.add_paragraph(sec.instructions)
            instr.runs[0].italic = True
            instr.paragraph_format.space_after = Pt(6)

        for item in sec.items:
            # Number label
            num_para = doc.add_paragraph()
            num_run = num_para.add_run(f"Number {item.number}.")
            num_run.bold = True
            num_run.font.size = Pt(11)
            num_para.paragraph_format.space_before = Pt(10)
            num_para.paragraph_format.space_after  = Pt(3)

            # Script block
            _add_shaded_script(doc, item.script)

            # Question
            q_para = doc.add_paragraph()
            q_run = q_para.add_run(f"{item.number}. {item.question}")
            q_run.bold = True
            q_para.paragraph_format.space_after = Pt(2)

            # Options
            for letter in ("A", "B", "C", "D"):
                opt_text = item.options.get(letter, "")
                if not opt_text:
                    continue
                opt_para = doc.add_paragraph()
                opt_para.paragraph_format.left_indent = Inches(0.25)
                opt_para.paragraph_format.space_after = Pt(1)
                opt_para.add_run(f"{letter}. {opt_text}")

            # Answer + explanation
            ans_para = doc.add_paragraph()
            ans_para.paragraph_format.space_before = Pt(4)
            ans_run = ans_para.add_run(f"【答案】{item.answer}")
            ans_run.bold = True
            ans_run.font.color.rgb = RGBColor(0x0D, 0x65, 0x2D)   # dark green
            if item.explanation:
                exp_run = ans_para.add_run(f"   {item.explanation}")
                exp_run.font.color.rgb = RGBColor(0x0D, 0x65, 0x2D)

        doc.add_paragraph()   # section spacer

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
